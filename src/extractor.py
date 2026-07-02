"""
extractor.py
Tahap 1: Text Extraction
Membaca dokumen input (.txt, .pdf, .docx) dan mengekstrak konten teks mentah.
"""

import os


def extract_text(file_path: str) -> str:
    """
    Mengekstrak teks dari file .txt, .pdf, atau .docx.

    Args:
        file_path: path menuju file yang akan diekstrak

    Returns:
        Teks mentah (raw string) hasil ekstraksi

    Raises:
        ValueError: jika format file tidak didukung
        FileNotFoundError: jika file tidak ditemukan
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File tidak ditemukan: {file_path}")

    ext = file_path.split(".")[-1].lower()

    if ext == "txt":
        return _extract_txt(file_path)
    elif ext == "pdf":
        return _extract_pdf(file_path)
    elif ext == "docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(
            f"Format file '.{ext}' tidak didukung. "
            "Gunakan .txt, .pdf, atau .docx"
        )


def _extract_txt(file_path: str) -> str:
    """Membaca file teks biasa."""
    with open(file_path, encoding="utf-8") as f:
        return f.read()


def _extract_pdf(file_path: str) -> str:
    """Mengekstrak teks dari PDF menggunakan PyMuPDF."""
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    try:
        pages = [page.get_text() for page in doc]
    finally:
        doc.close()
    return "\n".join(pages)


def _extract_docx(file_path: str) -> str:
    """Mengekstrak teks dari dokumen Word menggunakan python-docx."""
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Versi extract_text yang menerima bytes langsung (untuk upload Streamlit),
    tanpa perlu menyimpan file sementara ke disk terlebih dahulu kecuali
    untuk pdf/docx yang membutuhkan file-like object.

    Args:
        file_bytes: isi file dalam bentuk bytes
        filename: nama file asli (dipakai untuk menentukan ekstensi)

    Returns:
        Teks mentah hasil ekstraksi
    """
    import io

    ext = filename.split(".")[-1].lower()

    if ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")

    elif ext == "pdf":
        import fitz

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            pages = [page.get_text() for page in doc]
        finally:
            doc.close()
        return "\n".join(pages)

    elif ext == "docx":
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)

    else:
        raise ValueError(
            f"Format file '.{ext}' tidak didukung. "
            "Gunakan .txt, .pdf, atau .docx"
        )


if __name__ == "__main__":
    # Contoh penggunaan sederhana
    sample = "data/sample_docs/contoh_01.txt"
    if os.path.exists(sample):
        print(extract_text(sample))
    else:
        print("Jalankan dari root folder proyek dan pastikan file sample ada.")
